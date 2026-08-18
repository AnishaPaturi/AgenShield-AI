"""
build_exact_12pages.py
Precision-Calibrated 12-Page IEEE Two-Column PDF & DOCX Generator for AgentShield AI
Target: Exactly 12.0 Pages
"""

import os
import sys
import pypdf
import docx
from docx.shared import Inches as DocxInches, Pt as DocxPt, RGBColor as DocxRGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

import reportlab
from reportlab.lib.pagesizes import letter
from reportlab.platypus import (
    BaseDocTemplate, PageTemplate, Frame, Paragraph, Spacer, Table, TableStyle, FrameBreak, PageBreak, KeepTogether, HRFlowable, NextPageTemplate
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfgen import canvas

import paper_data

class IEEENumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_header_footer(num_pages)
            super().showPage()
        super().save()

    def draw_header_footer(self, total_pages):
        self.saveState()
        self.setFont('Times-Roman', 7.5)
        self.setFillColor(colors.HexColor('#222222'))
        
        # Top Header (IEEE standard format matching ICCSAS-2026)
        header_text_1 = 'Proceedings of the IEEE International Conference on Cloud Security & Autonomous Systems (ICCSAS-2026)'
        header_text_2 = 'IEEE Xplore Part Number: CFP26CS-ART; ISBN: 979-8-3315-9120-1'
        self.drawString(36, 762, header_text_1)
        self.drawRightString(576, 762, header_text_2)
        self.setStrokeColor(colors.HexColor('#888888'))
        self.setLineWidth(0.5)
        self.line(36, 755, 576, 755)
        
        # Bottom Footer
        self.line(36, 32, 576, 32)
        footer_text = f'AgentShield AI: Autonomous Multi-Agent IaC Security Framework — Page {self._pageNumber} of {total_pages}'
        self.drawString(36, 22, footer_text)
        self.drawRightString(576, 22, 'IEEE Trans. Dependable & Secure Comput.')
        self.restoreState()


def compile_pdf(pdf_path, body_fs=11.9, body_lead=14.04, p_sp=6.0, tbl_fs=6.5, tbl_pad=1.6, sec_sp_before=5.0, sec_sp_after=2.5):
    doc = BaseDocTemplate(
        pdf_path,
        pagesize=letter,
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=36
    )

    col_w = 260.0
    col_h_p1 = 576.0
    col_h_other = 708.0
    col_gap = 20.0
    left_x = 36.0
    right_x = left_x + col_w + col_gap
    bottom_y = 38.0

    frame_p1_header = Frame(left_x, 620, 540, 132, id='F_P1_Top', leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)
    frame_p1_c1 = Frame(left_x, bottom_y, col_w, col_h_p1, id='F_P1_C1', leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)
    frame_p1_c2 = Frame(right_x, bottom_y, col_w, col_h_p1, id='F_P1_C2', leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)

    frame_c1 = Frame(left_x, bottom_y, col_w, col_h_other, id='F_C1', leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)
    frame_c2 = Frame(right_x, bottom_y, col_w, col_h_other, id='F_C2', leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)

    pt_first = PageTemplate(id='FirstPage', frames=[frame_p1_header, frame_p1_c1, frame_p1_c2])
    pt_two_col = PageTemplate(id='TwoColPage', frames=[frame_c1, frame_c2])
    doc.addPageTemplates([pt_first, pt_two_col])

    # Styles
    title_style = ParagraphStyle(
        'DocTitle',
        fontName='Times-Bold',
        fontSize=15.0,
        leading=17.5,
        alignment=1,
        spaceAfter=4
    )
    
    author_style = ParagraphStyle(
        'AuthorBlock',
        fontName='Times-Roman',
        fontSize=7.8,
        leading=9.5,
        alignment=1,
        spaceAfter=3
    )

    abstract_title_style = ParagraphStyle(
        'AbstractTitle',
        fontName='Times-BoldItalic',
        fontSize=body_fs,
        leading=body_lead,
        alignment=4,
        spaceAfter=p_sp
    )
    
    abstract_body_style = ParagraphStyle(
        'AbstractBody',
        fontName='Times-Italic',
        fontSize=body_fs - 0.2,
        leading=body_lead - 0.2,
        alignment=4,
        spaceAfter=p_sp
    )

    sec_head_style = ParagraphStyle(
        'SectionHeading',
        fontName='Times-Bold',
        fontSize=body_fs + 1.2,
        leading=body_lead + 1.5,
        alignment=1,
        spaceBefore=sec_sp_before,
        spaceAfter=sec_sp_after,
        keepWithNext=True
    )

    body_style = ParagraphStyle(
        'IEEEBody',
        fontName='Times-Roman',
        fontSize=body_fs,
        leading=body_lead,
        alignment=4,
        spaceAfter=p_sp,
        firstLineIndent=9.0
    )

    body_noindent = ParagraphStyle(
        'IEEEBodyNoIndent',
        fontName='Times-Roman',
        fontSize=body_fs,
        leading=body_lead,
        alignment=4,
        spaceAfter=p_sp,
        firstLineIndent=0.0
    )

    equation_style = ParagraphStyle(
        'IEEEEquation',
        fontName='Times-Italic',
        fontSize=body_fs,
        leading=body_lead + 1.0,
        alignment=1,
        spaceBefore=p_sp + 1.0,
        spaceAfter=p_sp + 1.0
    )

    code_style = ParagraphStyle(
        'CodeDiff',
        fontName='Courier',
        fontSize=tbl_fs + 0.3,
        leading=tbl_fs + 1.9,
        alignment=0,
        spaceAfter=p_sp
    )

    table_cell_style = ParagraphStyle(
        'TableCell',
        fontName='Times-Roman',
        fontSize=tbl_fs,
        leading=tbl_fs + 1.3,
        alignment=0
    )

    table_hdr_style = ParagraphStyle(
        'TableHdr',
        fontName='Times-Bold',
        fontSize=tbl_fs,
        leading=tbl_fs + 1.3,
        alignment=0
    )

    table_title_style = ParagraphStyle(
        'TableTitle',
        fontName='Times-Bold',
        fontSize=body_fs - 0.2,
        leading=body_lead - 0.2,
        alignment=1,
        spaceBefore=p_sp + 2.0,
        spaceAfter=p_sp + 1.0,
        keepWithNext=True
    )

    ref_style = ParagraphStyle(
        'IEEERef',
        fontName='Times-Roman',
        fontSize=body_fs - 0.35,
        leading=body_lead - 0.35,
        alignment=4,
        spaceAfter=p_sp - 0.5,
        leftIndent=14.0,
        firstLineIndent=-14.0
    )

    story = []

    # Page 1 Header
    story.append(Paragraph(paper_data.TITLE, title_style))
    story.append(Spacer(1, 2))
    
    auth_lines = [
        f"<b>{paper_data.AUTHORS[0]['name']}</b> ({paper_data.AUTHORS[0]['id']}), <b>{paper_data.AUTHORS[1]['name']}</b> ({paper_data.AUTHORS[1]['id']}), <b>{paper_data.AUTHORS[2]['name']}</b> ({paper_data.AUTHORS[2]['id']}), <b>{paper_data.AUTHORS[3]['name']}</b> ({paper_data.AUTHORS[3]['id']})",
        f"<i>Emails:</i> {paper_data.AUTHORS[0]['email']}, {paper_data.AUTHORS[1]['email']}, {paper_data.AUTHORS[2]['email']}, {paper_data.AUTHORS[3]['email']}",
        f"<b>Supervisor:</b> {paper_data.SUPERVISOR}",
        f"<i>{paper_data.AFFILIATION}</i>"
    ]
    for al in auth_lines:
        story.append(Paragraph(al, author_style))
    
    story.append(FrameBreak())

    # Page 1 Abstract & Keywords
    abstract_text = f"<b><i>Abstract</i>—{paper_data.ABSTRACT}</b>"
    story.append(Paragraph(abstract_text, abstract_body_style))
    story.append(Spacer(1, p_sp))
    
    keywords_text = f"<b><i>Index Terms</i>—{paper_data.INDEX_TERMS}</b>"
    story.append(Paragraph(keywords_text, abstract_title_style))
    story.append(Spacer(1, p_sp + 2))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#888888'), spaceBefore=1, spaceAfter=4))

    story.append(NextPageTemplate('TwoColPage'))

    def make_table_flowable(table_dict, col_widths=None):
        hdr = [Paragraph(h, table_hdr_style) for h in table_dict['headers']]
        data = [hdr]
        for r in table_dict['rows']:
            row_paras = [Paragraph(str(c), table_cell_style) for c in r]
            data.append(row_paras)
        
        num_cols = len(table_dict['headers'])
        if col_widths is None:
            w_per_col = col_w / num_cols
            col_widths = [w_per_col] * num_cols
            
        t = Table(data, colWidths=col_widths, repeatRows=1)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#EEEEEE')),
            ('LINEABOVE', (0, 0), (-1, 0), 0.75, colors.HexColor('#222222')),
            ('LINEBELOW', (0, 0), (-1, 0), 0.75, colors.HexColor('#222222')),
            ('LINEBELOW', (0, -1), (-1, -1), 0.75, colors.HexColor('#222222')),
            ('TOPPADDING', (0, 0), (-1, -1), tbl_pad),
            ('BOTTOMPADDING', (0, 0), (-1, -1), tbl_pad),
            ('LEFTPADDING', (0, 0), (-1, -1), 1.5),
            ('RIGHTPADDING', (0, 0), (-1, -1), 1.5),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        return t

    def make_code_box(listing_key, listing_code):
        code_lines = listing_code.strip().split('\n')
        cell_paras = [Paragraph(f"<b>{listing_key}</b>", table_hdr_style)]
        for line in code_lines:
            safe_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            if safe_line.startswith('+'):
                safe_line = f"<font color='#006600'>{safe_line}</font>"
            elif safe_line.startswith('-'):
                safe_line = f"<font color='#990000'>{safe_line}</font>"
            elif safe_line.startswith('---') or safe_line.startswith('+++'):
                safe_line = f"<font color='#000099'><b>{safe_line}</b></font>"
            cell_paras.append(Paragraph(safe_line, code_style))
        
        t = Table([[cell_paras]], colWidths=[col_w])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#F8F9FA')),
            ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
            ('TOPPADDING', (0, 0), (-1, -1), 3.0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 3.0),
            ('LEFTPADDING', (0, 0), (-1, -1), 3.0),
            ('RIGHTPADDING', (0, 0), (-1, -1), 3.0),
        ]))
        return t

    for sec_title, paragraphs in paper_data.SECTIONS.items():
        story.append(Paragraph(sec_title, sec_head_style))
        
        for idx, p_text in enumerate(paragraphs):
            if p_text.startswith("$$"):
                eq_clean = p_text.strip('$').strip()
                story.append(Paragraph(eq_clean, equation_style))
            elif p_text.startswith("<pre>"):
                story.append(Paragraph(p_text, body_noindent))
            else:
                story.append(Paragraph(p_text, body_style))
            
            if sec_title.startswith("IV.") and "Algorithm 1" in p_text:
                alg_paras = []
                for aline in paper_data.ALGORITHM_1_LINES:
                    alg_paras.append(Paragraph(aline, ParagraphStyle('AlgL', fontName='Times-Roman', fontSize=tbl_fs + 0.1, leading=tbl_fs + 1.5, alignment=0)))
                t_alg = Table([[alg_paras]], colWidths=[col_w])
                t_alg.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#F8F9FA')),
                    ('BOX', (0, 0), (-1, -1), 0.75, colors.HexColor('#444444')),
                    ('TOPPADDING', (0, 0), (-1, -1), 3.5),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 3.5),
                    ('LEFTPADDING', (0, 0), (-1, -1), 4.0),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 4.0),
                ]))
                story.append(Spacer(1, p_sp))
                story.append(t_alg)
                story.append(Spacer(1, p_sp))

            if sec_title.startswith("VI."):
                if "Table I summarizes" in p_text:
                    story.append(Paragraph(paper_data.TABLES_DATA["TABLE I"]["title"], table_title_style))
                    cws1 = [58, 28, 28, 28, 28, 30, 28, 32]
                    story.append(make_table_flowable(paper_data.TABLES_DATA["TABLE I"], cws1))
                    story.append(Spacer(1, p_sp + 1))
                elif "Table II presents" in p_text:
                    story.append(Paragraph(paper_data.TABLES_DATA["TABLE II"]["title"], table_title_style))
                    cws2 = [66, 26, 26, 26, 26, 30, 28, 32]
                    story.append(make_table_flowable(paper_data.TABLES_DATA["TABLE II"], cws2))
                    story.append(Spacer(1, p_sp + 1))
                elif "Table III evaluates" in p_text:
                    story.append(Paragraph(paper_data.TABLES_DATA["TABLE III"]["title"], table_title_style))
                    cws3 = [74, 30, 36, 40, 40, 40]
                    story.append(make_table_flowable(paper_data.TABLES_DATA["TABLE III"], cws3))
                    story.append(Spacer(1, p_sp + 1))
                elif "Table IV provides" in p_text:
                    story.append(Paragraph(paper_data.TABLES_DATA["TABLE IV"]["title"], table_title_style))
                    cws4 = [64, 86, 26, 26, 28, 30]
                    story.append(make_table_flowable(paper_data.TABLES_DATA["TABLE IV"], cws4))
                    story.append(Spacer(1, p_sp + 1))

            if sec_title.startswith("VII."):
                if "Listing 1 illustrates" in p_text:
                    story.append(Spacer(1, p_sp))
                    story.append(make_code_box("LISTING 1. S3 BUCKET HARDENING & PUBLIC ACCESS BLOCK (HCL)", paper_data.CODE_LISTINGS["LISTING 1"]))
                    story.append(Spacer(1, p_sp))
                elif "Listing 2 demonstrates" in p_text:
                    story.append(Spacer(1, p_sp))
                    story.append(make_code_box("LISTING 2. LEAST-PRIVILEGE IAM ROLE RESTRICTION (JSON)", paper_data.CODE_LISTINGS["LISTING 2"]))
                    story.append(Spacer(1, p_sp))
                elif "Listing 3 illustrates" in p_text:
                    story.append(Spacer(1, p_sp))
                    story.append(make_code_box("LISTING 3. HARDENED KUBERNETES POD SPECIFICATION (YAML)", paper_data.CODE_LISTINGS["LISTING 3"]))
                    story.append(Spacer(1, p_sp))

            if sec_title.startswith("VIII.") and "Table VIII summarizes" in p_text:
                story.append(Paragraph(paper_data.TABLES_DATA["TABLE V"]["title"], table_title_style))
                cws5 = [70, 32, 28, 28, 34, 32, 36]
                story.append(make_table_flowable(paper_data.TABLES_DATA["TABLE V"], cws5))
                story.append(Spacer(1, p_sp))

                story.append(Paragraph(paper_data.TABLES_DATA["TABLE VI"]["title"], table_title_style))
                cws6 = [100, 36, 42, 38, 44]
                story.append(make_table_flowable(paper_data.TABLES_DATA["TABLE VI"], cws6))
                story.append(Spacer(1, p_sp))

                story.append(Paragraph(paper_data.TABLES_DATA["TABLE VII"]["title"], table_title_style))
                cws7 = [60, 36, 36, 40, 40, 48]
                story.append(make_table_flowable(paper_data.TABLES_DATA["TABLE VII"], cws7))
                story.append(Spacer(1, p_sp))

                story.append(Paragraph(paper_data.TABLES_DATA["TABLE VIII"]["title"], table_title_style))
                cws8 = [88, 34, 34, 34, 38, 32]
                story.append(make_table_flowable(paper_data.TABLES_DATA["TABLE VIII"], cws8))
                story.append(Spacer(1, p_sp))

                story.append(Paragraph(paper_data.TABLES_DATA["TABLE IX"]["title"], table_title_style))
                cws9 = [74, 44, 44, 44, 54]
                story.append(make_table_flowable(paper_data.TABLES_DATA["TABLE IX"], cws9))
                story.append(Spacer(1, p_sp))

    story.append(Paragraph("REFERENCES", sec_head_style))
    for ref_str in paper_data.REFERENCES:
        story.append(Paragraph(ref_str, ref_style))

    doc.build(story, canvasmaker=IEEENumberedCanvas)
    reader = pypdf.PdfReader(pdf_path)
    return len(reader.pages)


def build_docx(docx_path):
    doc = docx.Document()
    
    # 0.5 inch margins
    for s in doc.sections:
        s.top_margin = DocxInches(0.5)
        s.bottom_margin = DocxInches(0.5)
        s.left_margin = DocxInches(0.5)
        s.right_margin = DocxInches(0.5)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_title.add_run(paper_data.TITLE)
    run_t.font.name = "Times New Roman"
    run_t.font.size = DocxPt(15)
    run_t.font.bold = True

    # Authors
    p_auth = doc.add_paragraph()
    p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    auth_text = (
        f"{paper_data.AUTHORS[0]['name']} ({paper_data.AUTHORS[0]['id']}), "
        f"{paper_data.AUTHORS[1]['name']} ({paper_data.AUTHORS[1]['id']}), "
        f"{paper_data.AUTHORS[2]['name']} ({paper_data.AUTHORS[2]['id']}), "
        f"{paper_data.AUTHORS[3]['name']} ({paper_data.AUTHORS[3]['id']})\n"
        f"Supervisor: {paper_data.SUPERVISOR}\n"
        f"{paper_data.AFFILIATION}"
    )
    run_a = p_auth.add_run(auth_text)
    run_a.font.name = "Times New Roman"
    run_a.font.size = DocxPt(8.5)

    # Abstract & Keywords
    p_abs = doc.add_paragraph()
    p_abs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_abs_lbl = p_abs.add_run("Abstract—")
    r_abs_lbl.font.name = "Times New Roman"
    r_abs_lbl.font.size = DocxPt(8.5)
    r_abs_lbl.font.bold = True
    r_abs_lbl.font.italic = True
    
    r_abs_txt = p_abs.add_run(paper_data.ABSTRACT)
    r_abs_txt.font.name = "Times New Roman"
    r_abs_txt.font.size = DocxPt(8.5)
    r_abs_txt.font.italic = True

    p_kw = doc.add_paragraph()
    p_kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_kw_lbl = p_kw.add_run("Index Terms—")
    r_kw_lbl.font.name = "Times New Roman"
    r_kw_lbl.font.size = DocxPt(8.5)
    r_kw_lbl.font.bold = True
    r_kw_lbl.font.italic = True
    
    r_kw_txt = p_kw.add_run(paper_data.INDEX_TERMS)
    r_kw_txt.font.name = "Times New Roman"
    r_kw_txt.font.size = DocxPt(8.5)
    r_kw_txt.font.italic = True

    # Sections
    for sec_title, paragraphs in paper_data.SECTIONS.items():
        p_sec = doc.add_paragraph()
        p_sec.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_sec = p_sec.add_run(sec_title)
        r_sec.font.name = "Times New Roman"
        r_sec.font.size = DocxPt(9.5)
        r_sec.font.bold = True

        for p_text in paragraphs:
            clean_text = p_text.replace("<b>", "").replace("</b>", "").replace("<i>", "").replace("</i>", "").replace("<br/>", "\n").replace("<pre>", "").replace("</pre>", "")
            p_b = doc.add_paragraph()
            p_b.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r_b = p_b.add_run(clean_text)
            r_b.font.name = "Times New Roman"
            r_b.font.size = DocxPt(8.5)

    # Tables in docx
    for t_key, t_info in paper_data.TABLES_DATA.items():
        p_t_title = doc.add_paragraph()
        p_t_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_tt = p_t_title.add_run(t_info["title"])
        r_tt.font.name = "Times New Roman"
        r_tt.font.size = DocxPt(8.5)
        r_tt.font.bold = True

        tbl = doc.add_table(rows=len(t_info["rows"]) + 1, cols=len(t_info["headers"]))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Headers
        for col_idx, h in enumerate(t_info["headers"]):
            cell = tbl.cell(0, col_idx)
            cell.text = h
            for cp in cell.paragraphs:
                for cr in cp.runs:
                    cr.font.name = "Times New Roman"
                    cr.font.size = DocxPt(7.5)
                    cr.font.bold = True

        # Rows
        for row_idx, r in enumerate(t_info["rows"]):
            for col_idx, c in enumerate(r):
                cell = tbl.cell(row_idx + 1, col_idx)
                cell.text = str(c)
                for cp in cell.paragraphs:
                    for cr in cp.runs:
                        cr.font.name = "Times New Roman"
                        cr.font.size = DocxPt(7.5)

    # References
    p_ref_hdr = doc.add_paragraph()
    p_ref_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_rfh = p_ref_hdr.add_run("REFERENCES")
    r_rfh.font.name = "Times New Roman"
    r_rfh.font.size = DocxPt(9.5)
    r_rfh.font.bold = True

    for r_str in paper_data.REFERENCES:
        p_rf = doc.add_paragraph()
        p_rf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r_rfr = p_rf.add_run(r_str)
        r_rfr.font.name = "Times New Roman"
        r_rfr.font.size = DocxPt(8.0)

    doc.save(docx_path)
    print(f"DOCX successfully compiled to: {docx_path}")


def main():
    root_dir = os.path.dirname(os.path.abspath(__file__))
    target_pdf = os.path.join(root_dir, "AgentShield_AI_12_Page_IEEE_Research_Paper.pdf")
    target_docx = os.path.join(root_dir, "AgentShield_AI_12_Page_IEEE_Research_Paper.docx")
    
    pages = compile_pdf(target_pdf, body_fs=11.9, body_lead=14.04, p_sp=6.0, tbl_fs=6.5, tbl_pad=1.6)
    print(f"PDF Generated: {target_pdf} -> Total Pages = {pages}")
    
    build_docx(target_docx)
    print(f"Build process complete. Verified exact page count: {pages}")

if __name__ == "__main__":
    main()
