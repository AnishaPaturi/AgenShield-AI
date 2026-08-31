"""
build_6page_paper.py
Precision-Calibrated 6-Page IEEE Two-Column PDF & DOCX Generator for AgentShield AI
Includes Architecture Diagram (Fig. 1) and all Empirical Benchmark Figures (Fig. 2, Fig. 3, Fig. 4, Fig. 5)
Target: Exactly 6.0 Pages with Zero Text Overlap & Professional Publication Aesthetics
"""

import os
import sys
import re
import io
import pypdf
import docx
from docx.shared import Inches as DocxInches, Pt as DocxPt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

import reportlab
from reportlab.lib.pagesizes import letter
from reportlab.platypus import (
    BaseDocTemplate, PageTemplate, Frame, Paragraph, Spacer, Table, TableStyle, FrameBreak, HRFlowable, NextPageTemplate, Image as RLImage
)
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfgen import canvas

import matplotlib.pyplot as plt
from PIL import Image as PILImage

import paper_data_6pages as pdata

ROOT_DIR = os.path.dirname(os.path.abspath(__file__))
FIG_DIR = os.path.join(ROOT_DIR, "paper_figures")

IMAGE_FIG1 = os.path.join(ROOT_DIR, "image.png")
IMAGE_FIG2 = os.path.join(FIG_DIR, "fig_vulnerability_benchmark.png")
IMAGE_FIG3 = os.path.join(FIG_DIR, "fig_secret_and_remediation.png")
IMAGE_FIG4 = os.path.join(FIG_DIR, "fig_latency_breakdown.png")
IMAGE_FIG5 = os.path.join(FIG_DIR, "fig_ablation_and_impact.png")


def render_latex_flowable(latex_str, max_width=255.0, max_height=24.0):
    clean = latex_str.strip('$').strip()
    fig = plt.figure(figsize=(6.0, 0.55), dpi=300)
    fig.patch.set_alpha(0.0)
    ax = fig.add_axes([0, 0, 1, 1])
    ax.axis('off')
    ax.text(0.5, 0.5, f'${clean}$', fontsize=9.0, ha='center', va='center', color='#111111')
    
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=300, transparent=True, bbox_inches='tight', pad_inches=0.02)
    plt.close(fig)
    buf.seek(0)
    
    pil_im = PILImage.open(buf)
    w_px, h_px = pil_im.size
    aspect = h_px / w_px
    
    pt_w = min(max_width, (w_px * 72 / 300) * 0.75)
    pt_h = pt_w * aspect
    if pt_h > max_height:
        pt_h = max_height
        pt_w = pt_h / aspect
        
    buf.seek(0)
    return RLImage(buf, width=pt_w, height=pt_h)


class IEEENumberedCanvas6P(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_header_footer(num_pages)
            super().showPage()
        super().save()

    def draw_header_footer(self, total_pages):
        self.saveState()
        self.setFont('Times-Roman', 7.5)
        self.setFillColor(colors.HexColor('#222222'))
        self.setStrokeColor(colors.HexColor('#888888'))
        self.setLineWidth(0.5)
        
        # Bottom Footer
        self.line(36, 30, 576, 30)
        footer_text = f'AgentShield AI: Autonomous Multi-Agent IaC Security Framework - Page {self._pageNumber} of {total_pages}'
        self.drawString(36, 20, footer_text)
        self.drawRightString(576, 20, 'IEEE Trans. Dependable & Secure Comput.')
        self.restoreState()


def compile_pdf_6p(pdf_path, body_fs=10.3, body_lead=12.1, p_sp=4.2, tbl_fs=6.0, tbl_pad=1.3, sec_sp_before=3.8, sec_sp_after=1.6,
                   fig1_h=125, fig2_h=120, fig3_h=114, fig4_h=114, fig5_h=114):
    doc = BaseDocTemplate(
        pdf_path,
        pagesize=letter,
        leftMargin=36,
        rightMargin=36,
        topMargin=32,
        bottomMargin=32
    )

    col_w = 260.0
    col_h_p1 = 574.0
    col_h_other = 718.0
    col_gap = 20.0
    left_x = 36.0
    right_x = left_x + col_w + col_gap
    bottom_y = 34.0

    frame_p1_header = Frame(left_x, 610, 540, 144, id='F_P1_Top', leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)
    frame_p1_c1 = Frame(left_x, bottom_y, col_w, col_h_p1, id='F_P1_C1', leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)
    frame_p1_c2 = Frame(right_x, bottom_y, col_w, col_h_p1, id='F_P1_C2', leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)

    frame_c1 = Frame(left_x, bottom_y, col_w, col_h_other, id='F_C1', leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)
    frame_c2 = Frame(right_x, bottom_y, col_w, col_h_other, id='F_C2', leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)

    pt_first = PageTemplate(id='FirstPage', frames=[frame_p1_header, frame_p1_c1, frame_p1_c2])
    pt_two_col = PageTemplate(id='TwoColPage', frames=[frame_c1, frame_c2])
    doc.addPageTemplates([pt_first, pt_two_col])

    # Styles
    title_style = ParagraphStyle('DocTitle', fontName='Times-Bold', fontSize=13.5, leading=16.0, alignment=1, spaceAfter=3)
    auth_hdr_style = ParagraphStyle('AuthHdr', fontName='Times-Bold', fontSize=8.0, leading=9.5, alignment=1)
    auth_sub_style = ParagraphStyle('AuthSub', fontName='Times-Italic', fontSize=6.8, leading=8.0, alignment=1)
    auth_mail_style = ParagraphStyle('AuthMail', fontName='Times-Roman', fontSize=6.8, leading=8.0, alignment=1)
    abstract_title_style = ParagraphStyle('AbstractTitle', fontName='Times-BoldItalic', fontSize=body_fs, leading=body_lead, alignment=4, spaceAfter=p_sp)
    abstract_body_style = ParagraphStyle('AbstractBody', fontName='Times-Italic', fontSize=body_fs - 0.3, leading=body_lead - 0.2, alignment=4, spaceAfter=p_sp)
    sec_head_style = ParagraphStyle('SectionHeading', fontName='Times-Bold', fontSize=10.5, leading=12.5, alignment=1, spaceBefore=sec_sp_before, spaceAfter=sec_sp_after, keepWithNext=True)
    body_style = ParagraphStyle('IEEEBody', fontName='Times-Roman', fontSize=body_fs, leading=body_lead, alignment=4, spaceAfter=p_sp, firstLineIndent=7.0)
    body_noindent = ParagraphStyle('IEEEBodyNoIndent', fontName='Times-Roman', fontSize=body_fs, leading=body_lead, alignment=4, spaceAfter=p_sp, firstLineIndent=0.0)
    code_style = ParagraphStyle('CodeDiff', fontName='Courier', fontSize=tbl_fs + 0.1, leading=tbl_fs + 1.4, alignment=0, spaceAfter=p_sp)
    table_cell_style = ParagraphStyle('TableCell', fontName='Times-Roman', fontSize=tbl_fs, leading=tbl_fs + 1.1, alignment=0)
    table_hdr_style = ParagraphStyle('TableHdr', fontName='Times-Bold', fontSize=tbl_fs, leading=tbl_fs + 1.1, alignment=0)
    table_title_style = ParagraphStyle('TableTitle', fontName='Times-Bold', fontSize=body_fs - 0.3, leading=body_lead - 0.3, alignment=1, spaceBefore=p_sp, spaceAfter=p_sp - 1.0, keepWithNext=True)
    fig_caption_style = ParagraphStyle('FigCaption', fontName='Times-Italic', fontSize=tbl_fs + 0.3, leading=tbl_fs + 1.5, alignment=1, spaceBefore=2.0, spaceAfter=p_sp)
    ref_style = ParagraphStyle('IEEERef', fontName='Times-Roman', fontSize=body_fs - 0.5, leading=body_lead - 0.5, alignment=4, spaceAfter=p_sp - 1.2, leftIndent=10.0, firstLineIndent=-10.0)

    story = []

    # Page 1 Header (IEEE Standard Multi-Column Author Block)
    story.append(Paragraph(pdata.TITLE, title_style))
    story.append(Spacer(1, 2))
    
    col_cells = []
    for auth in pdata.AUTHORS:
        name = auth["name"]
        email = auth["email"]
        paras = [
            Paragraph(name, auth_hdr_style),
            Paragraph('<i>Department of Computer Science<br/>&amp; Engineering</i>', auth_sub_style),
            Paragraph('<i>Keshav Memorial Institute of Technology</i>', auth_sub_style),
            Paragraph('<i>Hyderabad, Telangana, India</i>', auth_sub_style),
            Paragraph(email, auth_mail_style)
        ]
        col_cells.append(paras)

    auth_table = Table([col_cells], colWidths=[135.0, 135.0, 135.0, 135.0])
    auth_table.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('LEFTPADDING', (0,0), (-1,-1), 1),
        ('RIGHTPADDING', (0,0), (-1,-1), 1),
        ('TOPPADDING', (0,0), (-1,-1), 1),
        ('BOTTOMPADDING', (0,0), (-1,-1), 1),
    ]))
    story.append(auth_table)
    story.append(FrameBreak())

    # Abstract & Keywords
    formatted_abs = pdata.ABSTRACT.replace('\n\n', '<br/><br/>')
    abstract_text = f"<b><i>Abstract</i>---{formatted_abs}</b>"
    story.append(Paragraph(abstract_text, abstract_body_style))
    story.append(Spacer(1, p_sp))
    keywords_text = f"<b><i>Index Terms</i>---{pdata.INDEX_TERMS}</b>"
    story.append(Paragraph(keywords_text, abstract_title_style))
    story.append(Spacer(1, p_sp))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#888888'), spaceBefore=1, spaceAfter=2))

    story.append(NextPageTemplate('TwoColPage'))

    def make_table_flowable(table_dict, col_widths=None):
        hdr = [Paragraph(h, table_hdr_style) for h in table_dict['headers']]
        data = [hdr]
        for r in table_dict['rows']:
            row_paras = [Paragraph(str(c), table_cell_style) for c in r]
            data.append(row_paras)
        num_cols = len(table_dict['headers'])
        if col_widths is None:
            w_per_col = col_w / num_cols
            col_widths = [w_per_col] * num_cols
        t = Table(data, colWidths=col_widths, repeatRows=1)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#EEEEEE')),
            ('LINEABOVE', (0, 0), (-1, 0), 0.6, colors.HexColor('#222222')),
            ('LINEBELOW', (0, 0), (-1, 0), 0.6, colors.HexColor('#222222')),
            ('LINEBELOW', (0, -1), (-1, -1), 0.6, colors.HexColor('#222222')),
            ('TOPPADDING', (0, 0), (-1, -1), tbl_pad),
            ('BOTTOMPADDING', (0, 0), (-1, -1), tbl_pad),
            ('LEFTPADDING', (0, 0), (-1, -1), 1.2),
            ('RIGHTPADDING', (0, 0), (-1, -1), 1.2),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        return t

    def make_code_box(listing_key, listing_code):
        code_lines = listing_code.strip().split('\n')
        cell_paras = [Paragraph(f"<b>{listing_key}</b>", table_hdr_style)]
        for line in code_lines:
            safe_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            if safe_line.startswith('+'):
                safe_line = f"<font color='#006600'>{safe_line}</font>"
            elif safe_line.startswith('-'):
                safe_line = f"<font color='#990000'>{safe_line}</font>"
            elif safe_line.startswith('---') or safe_line.startswith('+++'):
                safe_line = f"<font color='#000099'><b>{safe_line}</b></font>"
            cell_paras.append(Paragraph(safe_line, code_style))
        t = Table([[cell_paras]], colWidths=[col_w])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#F8F9FA')),
            ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
            ('TOPPADDING', (0, 0), (-1, -1), 2.0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 2.0),
            ('LEFTPADDING', (0, 0), (-1, -1), 2.0),
            ('RIGHTPADDING', (0, 0), (-1, -1), 2.0),
        ]))
        return t

    # SECTION I: Introduction
    story.append(Paragraph("I. Introduction", sec_head_style))
    for p in pdata.SECTIONS["I. Introduction"]:
        story.append(Paragraph(p, body_style))

    # SECTION II: Related Work
    story.append(Paragraph("II. Related Work", sec_head_style))
    for p in pdata.SECTIONS["II. Related Work"]:
        story.append(Paragraph(p, body_style))

    # SECTION III: System Architecture & Agent Methodology
    story.append(Paragraph("III. System Architecture & Agent Methodology", sec_head_style))
    sec3_paras = pdata.SECTIONS["III. System Architecture & Agent Methodology"]
    story.append(Paragraph(sec3_paras[0], body_style))
    
    # Insert Fig. 1 (Architecture Diagram)
    if os.path.exists(IMAGE_FIG1):
        story.append(Spacer(1, 1.5))
        story.append(RLImage(IMAGE_FIG1, width=col_w - 6, height=fig1_h))
        story.append(Spacer(1, 1.5))
        fig_cap1 = "<b>Fig. 1.</b> End-to-End System Architecture of AgentShield AI illustrating the 8-agent orchestration pipeline, Tree-sitter AST parsing, entropy-based secret scanning, hybrid RAG retrieval, dual-LLM consensus, and LocalStack sandbox validation."
        story.append(Paragraph(fig_cap1, fig_caption_style))
        story.append(Spacer(1, p_sp))

    for p in sec3_paras[1:]:
        story.append(Paragraph(p, body_style))

    # SECTION IV: Mathematical Formulation & Algorithmic Workflow
    story.append(Paragraph("IV. Mathematical Formulation & Algorithmic Workflow", sec_head_style))
    sec4_paras = pdata.SECTIONS["IV. Mathematical Formulation & Algorithmic Workflow"]
    for p in sec4_paras:
        if p.startswith("$$"):
            story.append(Spacer(1, 1.0))
            story.append(render_latex_flowable(p, max_width=col_w - 6))
            story.append(Spacer(1, 1.0))
        else:
            story.append(Paragraph(p, body_style))

    # Algorithm 1 Box
    alg_paras = []
    for aline in pdata.ALGORITHM_1_LINES_6P:
        alg_paras.append(Paragraph(aline, ParagraphStyle('AlgL', fontName='Times-Roman', fontSize=tbl_fs, leading=tbl_fs + 1.1, alignment=0)))
    t_alg = Table([[alg_paras]], colWidths=[col_w])
    t_alg.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#F8F9FA')),
        ('BOX', (0, 0), (-1, -1), 0.6, colors.HexColor('#444444')),
        ('TOPPADDING', (0, 0), (-1, -1), 2.0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2.0),
        ('LEFTPADDING', (0, 0), (-1, -1), 2.5),
        ('RIGHTPADDING', (0, 0), (-1, -1), 2.5),
    ]))
    story.append(Spacer(1, p_sp))
    story.append(t_alg)
    story.append(Spacer(1, p_sp))

    # SECTION V: Experimental Setup & Benchmark Methodology
    story.append(Paragraph("V. Experimental Setup & Benchmark Methodology", sec_head_style))
    for p in pdata.SECTIONS["V. Experimental Setup & Benchmark Methodology"]:
        story.append(Paragraph(p, body_style))

    # SECTION VI: Empirical Results & Discussion
    story.append(Paragraph("VI. Empirical Results & Discussion", sec_head_style))
    sec6_paras = pdata.SECTIONS["VI. Empirical Results & Discussion"]

    # Subsection VI-A
    story.append(Paragraph(sec6_paras[0], body_style))
    if os.path.exists(IMAGE_FIG2):
        story.append(Spacer(1, 1.5))
        story.append(RLImage(IMAGE_FIG2, width=col_w - 6, height=fig2_h))
        story.append(Spacer(1, 1.5))
        fig_cap2 = "<b>Fig. 2.</b> Comparative Vulnerability Detection Performance Across 2,450 Templates (Precision, Recall, F1-Score) highlighting AgentShield AI's 99.1% precision and 0.05% FPR."
        story.append(Paragraph(fig_cap2, fig_caption_style))
        story.append(Spacer(1, p_sp))

    story.append(Paragraph(pdata.TABLES_DATA_6P["TABLE I"]["title"], table_title_style))
    cws1 = [58, 28, 28, 28, 28, 30, 28, 32]
    story.append(make_table_flowable(pdata.TABLES_DATA_6P["TABLE I"], cws1))
    story.append(Spacer(1, p_sp))

    # Subsection VI-B
    story.append(Paragraph(sec6_paras[1], body_style))
    if os.path.exists(IMAGE_FIG3):
        story.append(Spacer(1, 1.5))
        story.append(RLImage(IMAGE_FIG3, width=col_w - 6, height=fig3_h))
        story.append(Spacer(1, 1.5))
        fig_cap3 = "<b>Fig. 3.</b> (a) Secret Detection Precision and False-Alarm Suppression; (b) Two-Tier LocalStack Sandbox Remediation Pass Rates (1st-Pass and Multi-Pass)."
        story.append(Paragraph(fig_cap3, fig_caption_style))
        story.append(Spacer(1, p_sp))

    story.append(Paragraph(pdata.TABLES_DATA_6P["TABLE II"]["title"], table_title_style))
    cws2 = [66, 26, 26, 26, 26, 30, 28, 32]
    story.append(make_table_flowable(pdata.TABLES_DATA_6P["TABLE II"], cws2))
    story.append(Spacer(1, p_sp))

    # Subsection VI-C
    story.append(Paragraph(sec6_paras[2], body_style))
    story.append(Paragraph(pdata.TABLES_DATA_6P["TABLE III"]["title"], table_title_style))
    cws3 = [74, 30, 36, 40, 40, 40]
    story.append(make_table_flowable(pdata.TABLES_DATA_6P["TABLE III"], cws3))
    story.append(Spacer(1, p_sp))

    # Subsection VI-D
    story.append(Paragraph(sec6_paras[3], body_style))
    if os.path.exists(IMAGE_FIG4):
        story.append(Spacer(1, 1.5))
        story.append(RLImage(IMAGE_FIG4, width=col_w - 6, height=fig4_h))
        story.append(Spacer(1, 1.5))
        fig_cap4 = "<b>Fig. 4.</b> Execution Latency Breakdown per Agent (Log Scale) across the 8-agent pipeline, demonstrating an average end-to-end runtime of 1.84s per IaC module."
        story.append(Paragraph(fig_cap4, fig_caption_style))
        story.append(Spacer(1, p_sp))

    story.append(Paragraph(pdata.TABLES_DATA_6P["TABLE IV"]["title"], table_title_style))
    cws4 = [80, 76, 34, 34, 36]
    story.append(make_table_flowable(pdata.TABLES_DATA_6P["TABLE IV"], cws4))
    story.append(Spacer(1, p_sp))

    # SECTION VII: Case Studies & Vulnerability Remediation
    story.append(Paragraph("VII. Case Studies & Vulnerability Remediation", sec_head_style))
    sec7_paras = pdata.SECTIONS["VII. Case Studies & Vulnerability Remediation"]
    story.append(Paragraph(sec7_paras[0], body_style))
    story.append(Spacer(1, p_sp - 1.0))
    story.append(make_code_box("Listing 1. S3 Bucket Hardening (Terraform HCL)", pdata.CODE_LISTINGS_6P["LISTING 1"]))
    story.append(Spacer(1, p_sp - 1.0))

    if len(sec7_paras) > 1:
        story.append(Paragraph(sec7_paras[1], body_style))
    story.append(Spacer(1, p_sp - 1.0))
    story.append(make_code_box("Listing 2. Least-Privilege IAM Scoping (JSON)", pdata.CODE_LISTINGS_6P["LISTING 2"]))
    story.append(Spacer(1, p_sp - 1.0))

    # SECTION VIII: Ablation Study & Cost Analysis
    story.append(Paragraph("VIII. Ablation Study & Cost Analysis", sec_head_style))
    sec8_paras = pdata.SECTIONS["VIII. Ablation Study & Cost Analysis"]
    story.append(Paragraph(sec8_paras[0], body_style))
    if os.path.exists(IMAGE_FIG5):
        story.append(Spacer(1, 1.5))
        story.append(RLImage(IMAGE_FIG5, width=col_w - 6, height=fig5_h))
        story.append(Spacer(1, 1.5))
        fig_cap5 = "<b>Fig. 5.</b> (a) Component Ablation Study across 500 templates; (b) Enterprise Cost and Mean Time to Remediate (MTTR) Reduction (99.99% decrease)."
        story.append(Paragraph(fig_cap5, fig_caption_style))
        story.append(Spacer(1, p_sp))

    story.append(Paragraph(pdata.TABLES_DATA_6P["TABLE V"]["title"], table_title_style))
    cws5 = [88, 34, 34, 34, 38, 32]
    story.append(make_table_flowable(pdata.TABLES_DATA_6P["TABLE V"], cws5))
    story.append(Spacer(1, p_sp))

    if len(sec8_paras) > 1:
        story.append(Paragraph(sec8_paras[1], body_style))
    story.append(Paragraph(pdata.TABLES_DATA_6P["TABLE VI"]["title"], table_title_style))
    cws6 = [74, 44, 44, 44, 54]
    story.append(make_table_flowable(pdata.TABLES_DATA_6P["TABLE VI"], cws6))
    story.append(Spacer(1, p_sp))

    # SECTION IX: Conclusion & Future Scope
    story.append(Paragraph("IX. Conclusion & Future Scope", sec_head_style))
    for p in pdata.SECTIONS["IX. Conclusion & Future Scope"]:
        story.append(Paragraph(p, body_style))

    # Reference Section
    story.append(Paragraph("Reference", sec_head_style))
    for ref_str in pdata.REFERENCES_6P:
        story.append(Paragraph(ref_str, ref_style))

    doc.build(story, canvasmaker=IEEENumberedCanvas6P)
    reader = pypdf.PdfReader(pdf_path)
    return len(reader.pages)


def sanitize_text(text):
    if not text:
        return ""
    t = re.sub(r'<[^>]+>', '', text)
    t = t.replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>').replace('&quot;', '"')
    t = "".join(c for c in t if ord(c) >= 32 or c in "\n\r\t")
    return t


def build_docx_6p(docx_path):
    doc = docx.Document()
    for s in doc.sections:
        s.top_margin = DocxInches(0.5)
        s.bottom_margin = DocxInches(0.5)
        s.left_margin = DocxInches(0.5)
        s.right_margin = DocxInches(0.5)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_title.add_run(pdata.TITLE)
    run_t.font.name = "Times New Roman"
    run_t.font.size = DocxPt(14)
    run_t.font.bold = True

    # 4-Column Author Table in DOCX matching IEEE layout
    tbl_auth = doc.add_table(rows=1, cols=len(pdata.AUTHORS))
    tbl_auth.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c_idx, auth in enumerate(pdata.AUTHORS):
        name = auth["name"]
        email = auth["email"]
        cell = tbl_auth.cell(0, c_idx)
        cell.width = DocxInches(1.85)
        
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = DocxPt(0)
        p.paragraph_format.space_after = DocxPt(1)
        r_name = p.add_run(name)
        r_name.font.name = "Times New Roman"
        r_name.font.size = DocxPt(8.5)
        r_name.font.bold = True
        
        p_sub = cell.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_before = DocxPt(0)
        p_sub.paragraph_format.space_after = DocxPt(1)
        r_sub = p_sub.add_run("Department of Computer Science\n& Engineering\nKeshav Memorial Institute of Technology\nHyderabad, Telangana, India")
        r_sub.font.name = "Times New Roman"
        r_sub.font.size = DocxPt(7.2)
        r_sub.font.italic = True
        
        p_mail = cell.add_paragraph()
        p_mail.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_mail.paragraph_format.space_before = DocxPt(0)
        p_mail.paragraph_format.space_after = DocxPt(2)
        r_mail = p_mail.add_run(email)
        r_mail.font.name = "Times New Roman"
        r_mail.font.size = DocxPt(7.2)

    abs_paras = pdata.ABSTRACT.split('\n\n')
    p_abs = doc.add_paragraph()
    p_abs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_abs_lbl = p_abs.add_run("Abstract---")
    r_abs_lbl.font.name = "Times New Roman"
    r_abs_lbl.font.size = DocxPt(8.0)
    r_abs_lbl.font.bold = True
    r_abs_lbl.font.italic = True
    
    r_abs_txt = p_abs.add_run(sanitize_text(abs_paras[0]))
    r_abs_txt.font.name = "Times New Roman"
    r_abs_txt.font.size = DocxPt(8.0)
    r_abs_txt.font.italic = True

    for ap in abs_paras[1:]:
        p_sub_abs = doc.add_paragraph()
        p_sub_abs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r_sub_abs = p_sub_abs.add_run(sanitize_text(ap))
        r_sub_abs.font.name = "Times New Roman"
        r_sub_abs.font.size = DocxPt(8.0)
        r_sub_abs.font.italic = True

    p_kw = doc.add_paragraph()
    p_kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_kw_lbl = p_kw.add_run("Index Terms---")
    r_kw_lbl.font.name = "Times New Roman"
    r_kw_lbl.font.size = DocxPt(8.0)
    r_kw_lbl.font.bold = True
    r_kw_lbl.font.italic = True
    
    r_kw_txt = p_kw.add_run(sanitize_text(pdata.INDEX_TERMS))
    r_kw_txt.font.name = "Times New Roman"
    r_kw_txt.font.size = DocxPt(8.0)
    r_kw_txt.font.italic = True

    def add_docx_fig(img_p, cap_t):
        if os.path.exists(img_p):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_picture(img_p, width=DocxInches(3.3))
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_cap = p_cap.add_run(cap_t)
            r_cap.font.name = "Times New Roman"
            r_cap.font.size = DocxPt(7.5)
            r_cap.font.italic = True

    # Helper for Docx sections
    for sec_title, paragraphs in pdata.SECTIONS.items():
        p_sec = doc.add_paragraph()
        p_sec.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_sec = p_sec.add_run(sec_title)
        r_sec.font.name = "Times New Roman"
        r_sec.font.size = DocxPt(11)
        r_sec.font.bold = True

        if sec_title.startswith("III."):
            add_docx_fig(IMAGE_FIG1, "Fig. 1. End-to-End System Architecture of AgentShield AI.")

        for p_text in paragraphs:
            if p_text.startswith("$$"):
                clean_text = p_text.strip('$').strip()
                p_b = doc.add_paragraph()
                p_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_b.paragraph_format.space_before = DocxPt(3)
                p_b.paragraph_format.space_after = DocxPt(3)
                r_b = p_b.add_run(clean_text)
                r_b.font.name = "Cambria Math"
                r_b.font.size = DocxPt(8.5)
                r_b.font.italic = True
            else:
                clean_text = sanitize_text(p_text)
                p_b = doc.add_paragraph()
                p_b.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                r_b = p_b.add_run(clean_text)
                r_b.font.name = "Times New Roman"
                r_b.font.size = DocxPt(8.0)

        if sec_title.startswith("VI."):
            add_docx_fig(IMAGE_FIG2, "Fig. 2. Comparative Vulnerability Detection Performance Across 2,450 Templates.")
            add_docx_fig(IMAGE_FIG3, "Fig. 3. Secret Detection & Sandbox Remediation Pass Rates.")
            add_docx_fig(IMAGE_FIG4, "Fig. 4. Execution Latency Breakdown per Agent (Total: 1.84s).")

        if sec_title.startswith("VIII."):
            add_docx_fig(IMAGE_FIG5, "Fig. 5. Component Ablation and Enterprise MTTR/Cost Reduction ROI.")

    for t_key, t_info in pdata.TABLES_DATA_6P.items():
        p_t_title = doc.add_paragraph()
        p_t_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_tt = p_t_title.add_run(t_info["title"])
        r_tt.font.name = "Times New Roman"
        r_tt.font.size = DocxPt(8.0)
        r_tt.font.bold = True

        tbl = doc.add_table(rows=len(t_info["rows"]) + 1, cols=len(t_info["headers"]))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for col_idx, h in enumerate(t_info["headers"]):
            cell = tbl.cell(0, col_idx)
            cell.text = sanitize_text(h)
            for cp in cell.paragraphs:
                for cr in cp.runs:
                    cr.font.name = "Times New Roman"
                    cr.font.size = DocxPt(7.0)
                    cr.font.bold = True

        for row_idx, r in enumerate(t_info["rows"]):
            for col_idx, c in enumerate(r):
                cell = tbl.cell(row_idx + 1, col_idx)
                cell.text = sanitize_text(str(c))
                for cp in cell.paragraphs:
                    for cr in cp.runs:
                        cr.font.name = "Times New Roman"
                        cr.font.size = DocxPt(7.0)

    p_ref_hdr = doc.add_paragraph()
    p_ref_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_rfh = p_ref_hdr.add_run("Reference")
    r_rfh.font.name = "Times New Roman"
    r_rfh.font.size = DocxPt(11)
    r_rfh.font.bold = True

    for r_str in pdata.REFERENCES_6P:
        p_rf = doc.add_paragraph()
        p_rf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r_rfr = p_rf.add_run(sanitize_text(r_str))
        r_rfr.font.name = "Times New Roman"
        r_rfr.font.size = DocxPt(7.2)

    try:
        doc.save(docx_path)
        print(f"DOCX 6-Page compiled: {docx_path}")
    except PermissionError:
        print(f"Notice: {docx_path} is currently open in Microsoft Word. Please close Word to allow overwrite.")


def calibrate_and_build():
    root_dir = os.path.dirname(os.path.abspath(__file__))
    target_pdf = os.path.join(root_dir, "AgentShield_AI_6_Page_IEEE_Research_Paper.pdf")
    target_docx = os.path.join(root_dir, "AgentShield_AI_6_Page_IEEE_Research_Paper.docx")

    print("Beginning precision calibration for exactly 6.0 pages...")
    
    # Targeted parameters for exactly 6.0 full pages
    params = {'body_fs': 10.4, 'body_lead': 12.2, 'p_sp': 4.4, 'tbl_fs': 6.0, 'tbl_pad': 1.4, 'fig1_h': 126, 'fig2_h': 122, 'fig3_h': 115, 'fig4_h': 115, 'fig5_h': 115}
    pages = compile_pdf_6p(target_pdf, **params)
    print(f"Generated PDF with {pages} pages using optimal parameters.")

    if pages != 6:
        # Fallback grid search
        for fs in [10.3, 10.2, 10.1, 10.0, 9.9]:
            lead = fs * 1.18
            pages = compile_pdf_6p(target_pdf, body_fs=fs, body_lead=lead, p_sp=4.2, tbl_fs=6.0, tbl_pad=1.3, fig1_h=124, fig2_h=120, fig3_h=114, fig4_h=114, fig5_h=114)
            print(f"Fallback grid fs={fs}: pages={pages}")
            if pages == 6:
                break

    build_docx_6p(target_docx)
    print(f"Completed! PDF: {target_pdf} ({pages} pages) | DOCX: {target_docx}")


def main():
    calibrate_and_build()


if __name__ == "__main__":
    main()
